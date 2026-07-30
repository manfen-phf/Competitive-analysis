from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

OUT = r"D:\ChatGPT\Projects\玉林商数据汇总\广西外卖竞争态势分析工具-第一版设计说明.docx"

BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F2F4F7"
MUTED = "666666"
BLACK = "1A1A1A"


def set_run_font(run, size=None, color=None, bold=None, italic=None):
    run.font.name = "Microsoft YaHei"
    run._element.rPr.rFonts.set(qn("w:ascii"), "Microsoft YaHei")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Microsoft YaHei")
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    if size:
        run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for side, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_widths(table, widths):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    tbl_w.set(qn("w:w"), "9360")
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = OxmlElement("w:tblInd")
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")
    tbl_pr.append(tbl_ind)
    width_dxa = [int(round(width.inches * 1440)) for width in widths]
    width_dxa[-1] += 9360 - sum(width_dxa)
    grid_cols = table._tbl.tblGrid.gridCol_lst
    for grid_col, value in zip(grid_cols, width_dxa):
        grid_col.set(qn("w:w"), str(value))
    for row in table.rows:
        for cell, value in zip(row.cells, width_dxa):
            cell.width = Inches(value / 1440)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.first_child_found_in("w:tcW")
            tc_w.set(qn("w:w"), str(value))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_page_field(paragraph):
    run = paragraph.add_run()
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = " PAGE "
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char1)
    run._r.append(instr_text)
    run._r.append(fld_char2)


def add_heading(doc, text, level=1):
    p = doc.add_paragraph(style=f"Heading {level}")
    p.add_run(text)
    return p


def add_body(doc, text):
    p = doc.add_paragraph(style="Normal")
    p.add_run(text)
    return p


def add_numbered(doc, text):
    p = doc.add_paragraph(style="List Number")
    p.add_run(text)
    return p


doc = Document()
section = doc.sections[0]
section.top_margin = Inches(1)
section.bottom_margin = Inches(1)
section.left_margin = Inches(1)
section.right_margin = Inches(1)
section.header_distance = Inches(0.492)
section.footer_distance = Inches(0.492)

styles = doc.styles
normal = styles["Normal"]
normal.font.name = "Microsoft YaHei"
normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
normal.font.size = Pt(11)
normal.paragraph_format.space_after = Pt(6)
normal.paragraph_format.line_spacing = 1.10

for name, size, color, before, after in [
    ("Heading 1", 16, BLUE, 16, 8),
    ("Heading 2", 13, BLUE, 12, 6),
    ("Heading 3", 12, DARK_BLUE, 8, 4),
]:
    style = styles[name]
    style.font.name = "Microsoft YaHei"
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    style.font.size = Pt(size)
    style.font.color.rgb = RGBColor.from_string(color)
    style.paragraph_format.space_before = Pt(before)
    style.paragraph_format.space_after = Pt(after)

for name in ("List Number",):
    style = styles[name]
    style.font.name = "Microsoft YaHei"
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    style.font.size = Pt(11)
    style.paragraph_format.space_after = Pt(8)
    style.paragraph_format.line_spacing = 1.167

# Quiet running header and footer.
header_p = section.header.paragraphs[0]
header_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
header_run = header_p.add_run("广西外卖竞争态势分析工具 | 第一版设计说明")
set_run_font(header_run, size=8.5, color=MUTED)
footer_p = section.footer.paragraphs[0]
footer_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
footer_run = footer_p.add_run("内部评审稿 | 第 ")
set_run_font(footer_run, size=8.5, color=MUTED)
add_page_field(footer_p)
tail = footer_p.add_run(" 页")
set_run_font(tail, size=8.5, color=MUTED)

# Memo masthead title block.
title = doc.add_paragraph()
title.paragraph_format.space_before = Pt(8)
title.paragraph_format.space_after = Pt(4)
run = title.add_run("第一版产品设计说明")
set_run_font(run, size=23, color=BLACK, bold=True)

subtitle = doc.add_paragraph()
subtitle.paragraph_format.space_after = Pt(14)
run = subtitle.add_run("商家维度外卖竞争态势分析工具")
set_run_font(run, size=14, color=MUTED)

metadata = doc.add_table(rows=4, cols=2)
set_table_widths(metadata, [Inches(1.2), Inches(5.3)])
for idx, (label, value) in enumerate([
    ("版本", "V1.0"),
    ("用途", "内部需求评审与开发规划"),
    ("核心对象", "一线城市运营人员、总部/管理层、主数据维护人"),
    ("状态", "已确认设计，待进入开发计划"),
]):
    left, right = metadata.rows[idx].cells
    set_cell_shading(left, LIGHT_GRAY)
    p = left.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    set_run_font(p.add_run(label), size=10.5, color=BLACK, bold=True)
    p = right.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    set_run_font(p.add_run(value), size=10.5, color=BLACK)

add_heading(doc, "1. 目标与使用者")
add_body(doc, "本工具用于比较同一商家在美团和 B家（产品内对饿了么的统一称呼）的竞争价格力。系统服务于截图采集、严格识别入库和管理层分析三类工作。")

roles = doc.add_table(rows=1, cols=2)
set_table_widths(roles, [Inches(1.8), Inches(4.7)])
for cell, text in zip(roles.rows[0].cells, ["角色", "主要职责"]):
    set_cell_shading(cell, LIGHT_BLUE)
    p = cell.paragraphs[0]; p.paragraph_format.space_after = Pt(0)
    set_run_font(p.add_run(text), 10.5, BLACK, True)
for role, duty in [
    ("一线城市运营", "选择城市与商家后上传订单详情截图。"),
    ("总部/管理层", "按时间、城市、BD、商家查看价格构成、差异与排名。"),
    ("主数据维护人", "仅由产品负责人导入商家与 BD 主数据。"),
]:
    cells = roles.add_row().cells
    for cell, text in zip(cells, [role, duty]):
        p = cell.paragraphs[0]; p.paragraph_format.space_after = Pt(0)
        set_run_font(p.add_run(text), 10.5, BLACK)

add_heading(doc, "2. 主数据与归属规则")
add_body(doc, "商家主数据以 Excel 导入，商家 ID 为主键。一个商家可存在多段 BD 归属记录，系统以截图上传日期匹配当时有效的 BD 归属。")

master = doc.add_table(rows=1, cols=2)
set_table_widths(master, [Inches(2.2), Inches(4.3)])
for cell, text in zip(master.rows[0].cells, ["必填列", "说明"]):
    set_cell_shading(cell, LIGHT_BLUE)
    p = cell.paragraphs[0]; p.paragraph_format.space_after = Pt(0)
    set_run_font(p.add_run(text), 10.5, BLACK, True)
for field, note in [
    ("商家 ID", "唯一主键，用于商家搜索、关联与去重。"),
    ("商家名称", "在选择城市后支持按名称搜索。"),
    ("城市", "限制可选商家范围。"),
    ("BD 姓名", "按有效期自动回填至上传记录。"),
    ("生效开始日 / 结束日", "支持 BD 历史归属回溯。"),
]:
    cells = master.add_row().cells
    for cell, text in zip(cells, [field, note]):
        p = cell.paragraphs[0]; p.paragraph_format.space_after = Pt(0)
        set_run_font(p.add_run(text), 10.5, BLACK)

add_body(doc, "普通使用不需要登录。主数据导入使用独立入口并以管理员口令保护；导入前展示新增、更新、失效与错误预览，且仅在整份文件校验通过后一次性生效。")

add_heading(doc, "3. 上传、识别与严格校验")
add_heading(doc, "3.1 操作流程", level=2)
for step in [
    "选择城市。",
    "在所选城市内按商家 ID 或商家名称搜索并选择商家。",
    "上传订单详情截图；每张截图只对应一个商家和一个平台。",
    "AI 从图片中识别平台（美团或 B家）及全部必填金额字段。",
    "系统执行字段完整性、格式、金额关系、置信度及重复记录校验。",
    "通过校验的数据入库；失败图片不入库，仅返回明确失败原因。",
]:
    add_numbered(doc, step)

add_heading(doc, "3.2 必填识别字段", level=2)
fields = [
    "菜品原价", "餐盒费", "平台红包", "原价配送费", "减配送费", "实付配送费",
    "商家结算金额", "用户实付", "其他活动", "技术服务费", "配送服务费", "商家费率",
]
field_table = doc.add_table(rows=1, cols=3)
set_table_widths(field_table, [Inches(2.17), Inches(2.17), Inches(2.16)])
for cell, text in zip(field_table.rows[0].cells, ["字段组 1", "字段组 2", "字段组 3"]):
    set_cell_shading(cell, LIGHT_BLUE)
    p = cell.paragraphs[0]; p.paragraph_format.space_after = Pt(0)
    set_run_font(p.add_run(text), 10.5, BLACK, True)
for row in range(4):
    cells = field_table.add_row().cells
    for cell, text in zip(cells, fields[row*3:(row+1)*3]):
        p = cell.paragraphs[0]; p.paragraph_format.space_after = Pt(0)
        set_run_font(p.add_run(text), 10.5, BLACK)

add_body(doc, "任一字段缺失、平台无法识别、置信度不足或金额关系不合理时，整张截图上传失败。系统不提供人工补录或编辑。商家费率还会通过技术服务费与计费基数做合理性校验；图片指纹与订单号双重去重。")

add_heading(doc, "4. 时间口径与筛选")
add_body(doc, "报表日期一律采用上传时间，不使用截图中的下单或完成时间。日期筛选支持日、周、月、年和范围选择，采用双面板日历，不提供“最近 7 天”“本月”等快捷选择。")
add_body(doc, "周按每年内自然周编号：1 月 1 日至该周周日为 W1，即使不足七天；随后每周一至周日为完整周。例如，2026-01-01 至 2026-01-04 为 W1，2026-01-05 至 2026-01-11 为 W2。")

add_heading(doc, "5. 竞争价格力看板")
add_body(doc, "所有金额按商家、平台和筛选周期汇总后除以有效订单数，统一呈现为每单平均值。仅统计通过识别、校验和去重的记录，并同时展示有效订单数。")

dashboard = doc.add_table(rows=1, cols=2)
set_table_widths(dashboard, [Inches(2.0), Inches(4.5)])
for cell, text in zip(dashboard.rows[0].cells, ["模块", "设计内容"]):
    set_cell_shading(cell, LIGHT_BLUE)
    p = cell.paragraphs[0]; p.paragraph_format.space_after = Pt(0)
    set_run_font(p.add_run(text), 10.5, BLACK, True)
for module, detail in [
    ("全局筛选", "日、周、月、年与日期范围；城市、BD、商家。平台固定并列展示美团与 B家。"),
    ("核心概览", "订单数、用户实付、平台红包、配送费、商家结算金额及平台差异。"),
    ("价格构成", "菜品原价、餐盒费、平台红包、其他活动、配送费、技术/配送服务费及商家费率对比。"),
    ("商家排名", "按用户实付差异或商家结算差异排序，识别价格优势平台。"),
    ("下钻与追溯", "从城市或 BD 汇总进入商家；查看成功截图的识别字段与原图，不可修改。"),
]:
    cells = dashboard.add_row().cells
    for cell, text in zip(cells, [module, detail]):
        p = cell.paragraphs[0]; p.paragraph_format.space_after = Pt(0)
        set_run_font(p.add_run(text), 10.5, BLACK)

add_heading(doc, "6. 次级采集健康页面")
add_body(doc, "采集健康不抢占价格力主看板，单独呈现有效截图数、失败数、重复数、商家覆盖率、按城市和 BD 的上传情况，并提供失败原因与原图追溯。")

add_heading(doc, "7. 技术方案与验收")
add_body(doc, "系统采用定制 Web 前端、服务端 API、关系型数据库、图片对象存储和多模态 AI 识别服务。主数据、导入版本、图片上传、识别结果、校验结果和分析汇总分开存储。前端不直接写入业务数据。")
add_body(doc, "验收时，每个平台至少准备字段齐全、缺字段、金额异常、重复订单、BD 已变更五类截图，验证严格拦截、去重、历史归属和报表口径。没有数据时，图表必须显示清晰的零状态。")

add_heading(doc, "8. 第一版范围边界")
add_body(doc, "第一版不提供普通账号、人工改数、人工补字段或第三方平台数据直连。自动采集、权限体系、更多平台和预警规则将在本版数据模型稳定后另行规划。")

doc.core_properties.title = "广西外卖竞争态势分析工具 - 第一版设计说明"
doc.core_properties.subject = "内部需求评审与开发规划"
doc.core_properties.author = "Codex"
doc.save(OUT)
print(OUT)
